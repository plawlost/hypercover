import os
import asyncio
import logging
import zipfile
from datetime import datetime, date
from typing import List, Dict, Optional, Callable
import pandas as pd
from groq import AsyncGroq
from pathlib import Path
import aiofiles
from docx import Document
import pypandoc
from io import BytesIO
import json
from functools import lru_cache
from concurrent.futures import ThreadPoolExecutor
import aiodns
from charset_normalizer import detect
from tenacity import retry, stop_after_attempt, wait_exponential
from prometheus_client import Counter, Histogram
from openai import AsyncOpenAI
from bs4 import BeautifulSoup
import aiohttp
from redis import asyncio as aioredis
from templates.cover_letter_templates import get_template_info, get_template_prompt

# Performance metrics
GENERATION_TIME = Histogram('cover_letter_generation_seconds', 'Time spent generating cover letters')
CACHE_HITS = Counter('cache_hits_total', 'Total number of cache hits')
API_ERRORS = Counter('api_errors_total', 'Total number of API errors')

class BulkCoverLetterGenerator:
    def __init__(self, groq_api_key: str, deepseek_api_key: str):
        self.groq_client = AsyncGroq(api_key=groq_api_key)
        self.openai_client = AsyncOpenAI(
            api_key=deepseek_api_key,
            base_url="https://api.deepseek.com"
        )
        # Create output directory with parents
        self.output_dir = Path("generated_letters")
        try:
            self.output_dir.mkdir(parents=True, exist_ok=True)
            if not self.output_dir.exists():
                raise Exception("Failed to create output directory")
        except Exception as e:
            raise Exception(f"Error creating output directory: {str(e)}")
            
        self.company_cache = {}
        self.session = None
        self.executor = ThreadPoolExecutor(max_workers=os.cpu_count() * 2)
        self.dns_resolver = None
        self.redis_client = None
        self.template_cache = {}
        self.logger = logging.getLogger(__name__)
        self.event_loop = None
        
        # Configure optimal chunk sizes based on testing
        self.COMPANY_BATCH_SIZE = 25
        self.LETTER_BATCH_SIZE = 20
        self.MAX_CONCURRENT_REQUESTS = 50
        
    async def initialize(self):
        """Initialize async resources with optimized settings"""
        try:
            # Configure logging
            logging.basicConfig(level=logging.INFO)
            
            # Initialize event loop
            try:
                self.event_loop = asyncio.get_event_loop()
            except RuntimeError:
                self.event_loop = asyncio.new_event_loop()
                asyncio.set_event_loop(self.event_loop)
            
            # Optimize HTTP session
            timeout = aiohttp.ClientTimeout(total=30, connect=5, sock_connect=5, sock_read=10)
            connector = aiohttp.TCPConnector(
                limit=self.MAX_CONCURRENT_REQUESTS,
                ttl_dns_cache=300,
                use_dns_cache=True,
                ssl=False,
                keepalive_timeout=60
            )
            self.session = aiohttp.ClientSession(
                timeout=timeout,
                connector=connector,
                headers={'User-Agent': 'Mozilla/5.0'}
            )
            
            # Initialize DNS resolver
            self.dns_resolver = aiodns.DNSResolver(loop=self.event_loop)
            
            # Initialize Redis with proper error handling
            self.redis_client = None
            try:
                redis = await aioredis.Redis.from_url(
                    'redis://localhost',
                    max_connections=20,
                    socket_timeout=1,
                    retry_on_timeout=True,
                    health_check_interval=30
                )
                # Validate connection with a ping and set operation
                if await redis.ping() and await redis.set('test_key', 'test_value', ex=1):
                    self.redis_client = redis
                    self.logger.info("Redis connection established and validated")
                else:
                    self.logger.warning("Redis connection failed validation checks")
                    await redis.close()
            except Exception as e:
                self.logger.warning(f"Redis not available, falling back to in-memory cache: {str(e)}")
                if 'redis' in locals():
                    await redis.close()
            
            # Warm up template cache
            await self._preload_templates()
            
        except Exception as e:
            self.logger.error(f"Error in initialization: {str(e)}")
            raise

    async def _is_redis_available(self) -> bool:
        """Check if Redis is available and connected"""
        if not self.redis_client:
            return False
        try:
            return await self.redis_client.ping()
        except Exception:
            return False

    async def _safe_redis_operation(self, operation):
        """Safely execute a Redis operation with connection check"""
        if not await self._is_redis_available():
            return None
        try:
            return await operation()
        except Exception as e:
            self.logger.warning(f"Redis operation failed: {str(e)}")
            return None

    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    async def _fetch_with_retry(self, url: str) -> str:
        """Fetch URL with exponential backoff retry"""
        try:
            async with self.session.get(url) as response:
                response.raise_for_status()
                return await response.text()
        except Exception as e:
            API_ERRORS.inc()
            self.logger.error(f"Error fetching {url}: {str(e)}")
            raise

    async def _preload_templates(self):
        """Preload and cache templates for faster access"""
        templates = get_template_info()
        for template_id in templates.keys():
            template = get_template_prompt(template_id)
            self.template_cache[template_id] = template
            
            if self.redis_client:
                await self._safe_redis_operation(
                    lambda: self.redis_client.set(
                        f"template:{template_id}",
                        json.dumps(template)
                    )
                )

    @GENERATION_TIME.time()
    async def process_spreadsheet(self, spreadsheet_path: str, user_profile: Dict, template_id: str, formats: List[str] = None, progress_callback=None):
        """Process spreadsheet with optimized batch processing"""
        try:
            # Default to DOCX if no formats specified
            if not formats:
                formats = ['docx']
                
            df = pd.read_csv(spreadsheet_path)
            total_tasks = len(df)
            completed_tasks = 0
            total_progress = 0  # Track overall progress including company info and letter generation
            
            # Create progress update function that handles both async and sync callbacks
            async def update_progress(progress_value):
                nonlocal total_progress
                total_progress = progress_value
                if progress_callback:
                    try:
                        if asyncio.iscoroutinefunction(progress_callback):
                            await progress_callback(progress_value)
                        else:
                            progress_callback(progress_value)
                    except Exception as e:
                        self.logger.error(f"Error updating progress: {str(e)}")
            
            # Pre-fetch and cache company data in optimal batches
            company_names = df['company_name'].unique()
            company_batches = [
                company_names[i:i + self.COMPANY_BATCH_SIZE] 
                for i in range(0, len(company_names), self.COMPANY_BATCH_SIZE)
            ]
            
            # Use semaphore to control concurrent API requests
            sem = asyncio.Semaphore(min(self.MAX_CONCURRENT_REQUESTS, 10))  # Limit concurrent connections
            
            async def fetch_with_semaphore(name):
                async with sem:
                    try:
                        # Ensure we have a valid event loop
                        if not self.event_loop or self.event_loop.is_closed():
                            self.event_loop = asyncio.get_event_loop()
                        return await self.gather_company_info(name)
                    except Exception as e:
                        self.logger.error(f"Error fetching company info for {name}: {str(e)}")
                        return {"name": name, "description": f"Company information unavailable"}
            
            # Fetch company data with controlled concurrency (40% of total progress)
            company_info_cache = {}
            for i, batch in enumerate(company_batches):
                company_tasks = [fetch_with_semaphore(name) for name in batch]
                batch_results = await asyncio.gather(*company_tasks, return_exceptions=True)
                
                # Cache company info, handling exceptions
                for name, result in zip(batch, batch_results):
                    if isinstance(result, Exception):
                        self.logger.error(f"Error processing company {name}: {str(result)}")
                        company_info_cache[name] = {
                            "name": name,
                            "description": "Company information unavailable"
                        }
                    else:
                        company_info_cache[name] = result
                
                # Update progress (company info is 40% of total progress)
                progress = (i + 1) / len(company_batches) * 40
                await update_progress(progress)
            
            # Process cover letters in optimized batches (remaining 60% of progress)
            results = []
            letter_batches = [
                df.iloc[i:i + self.LETTER_BATCH_SIZE] 
                for i in range(0, len(df), self.LETTER_BATCH_SIZE)
            ]
            
            for i, batch_df in enumerate(letter_batches):
                batch_tasks = []
                
                for _, row in batch_df.iterrows():
                    company_name = row['company_name']
                    company_info = company_info_cache.get(company_name, {
                        "name": company_name,
                        "description": "Company information unavailable"
                    })
                    
                    task = self.generate_single_letter(
                        company_name=company_name,
                        position=row['position'],
                        user_profile=user_profile,
                        template_id=template_id,
                        formats=formats,
                        notes=row.get('notes', ''),
                        company_info=company_info
                    )
                    batch_tasks.append(task)
                
                # Process batch with error handling
                batch_results = await asyncio.gather(*batch_tasks, return_exceptions=True)
                
                # Handle results
                for result in batch_results:
                    if isinstance(result, Exception):
                        self.logger.error(f"Error in letter generation: {str(result)}")
                        API_ERRORS.inc()
                    else:
                        results.append(result)
                
                # Update progress (letter generation is 60% of total progress)
                progress = 40 + ((i + 1) / len(letter_batches) * 60)
                await update_progress(progress)
            
            # Ensure final progress update
            await update_progress(100)
            
            # Package results
            try:
                return await self._package_results(results)
            except Exception as e:
                self.logger.error(f"Error packaging results: {str(e)}")
                raise
            
        except Exception as e:
            self.logger.error(f"Error in bulk processing: {str(e)}")
            raise
        finally:
            # Clean up any temporary files
            try:
                if os.path.exists(spreadsheet_path):
                    os.remove(spreadsheet_path)
            except Exception as e:
                self.logger.warning(f"Error cleaning up temporary file: {str(e)}")

    async def _async_remove(self, path: str):
        """Asynchronously remove a file"""
        try:
            await asyncio.get_event_loop().run_in_executor(
                self.executor, os.remove, path
            )
        except:
            pass

    @lru_cache(maxsize=1000)
    async def gather_company_info(self, company_name: str) -> Dict:
        """Gather and cache company information with optimized parallel fetching"""
        # Check Redis cache first
        if self.redis_client:
            cached_data = await self.redis_client.get(f"company:{company_name}")
            if cached_data:
                return json.loads(cached_data)
        
        if company_name in self.company_cache:
            return self.company_cache[company_name]
            
        try:
            # Fetch all data in parallel with timeouts
            tasks = [
                asyncio.create_task(self._fetch_linkedin_data(company_name)),
                asyncio.create_task(self._fetch_website_data(company_name)),
                asyncio.create_task(self._fetch_glassdoor_data(company_name)),
                asyncio.create_task(self._fetch_news_data(company_name))
            ]
            
            # Wait for all tasks with a timeout
            done, pending = await asyncio.wait(
                tasks,
                timeout=5,  # Aggressive timeout
                return_when=asyncio.ALL_COMPLETED
            )
            
            # Cancel any pending tasks
            for task in pending:
                task.cancel()
            
            # Get results, defaulting to empty data for failed tasks
            results = []
            for task in done:
                try:
                    results.append(task.result())
                except:
                    results.append({})
            
            linkedin_data, website_data, glassdoor_data, news_data = (
                results + [{}] * (4 - len(results))
            )[:4]
            
            # Process data in parallel
            keywords_task = asyncio.create_task(self._extract_keywords(website_data))
            culture_task = asyncio.create_task(self._analyze_company_culture(
                linkedin_data.get("description", ""),
                glassdoor_data.get("culture", ""),
                website_data
            ))
            tech_stack_task = asyncio.create_task(self._extract_tech_stack(
                website_data,
                linkedin_data.get("description", "")
            ))
            
            # Wait for processing with timeout
            website_keywords, culture_analysis, tech_stack = await asyncio.gather(
                keywords_task, culture_task, tech_stack_task,
                return_exceptions=True
            )
            
            # Handle exceptions
            if isinstance(website_keywords, Exception): website_keywords = []
            if isinstance(culture_analysis, Exception): culture_analysis = {}
            if isinstance(tech_stack, Exception): tech_stack = []
            
            company_data = {
                "name": company_name,
                "description": linkedin_data.get("description", ""),
                "industry": linkedin_data.get("industry", ""),
                "website_data": website_data,
                "culture": {
                    "values": culture_analysis.get("values", []),
                    "work_environment": culture_analysis.get("environment", ""),
                    "benefits": glassdoor_data.get("benefits", [])
                },
                "recent_news": news_data.get("articles", [])[:3],
                "keywords": website_keywords,
                "technologies": tech_stack
            }
            
            # Cache the results
            self.company_cache[company_name] = company_data
            if self.redis_client:
                await self.redis_client.set(
                    f"company:{company_name}",
                    json.dumps(company_data),
                    expire=3600  # Cache for 1 hour
                )
            
            return company_data
            
        except Exception as e:
            return {"name": company_name, "description": f"Company information for {company_name}"}

    async def _fetch_linkedin_data(self, company_name: str) -> Dict:
        """Fetch LinkedIn data with optimized retry logic and parsing"""
        retry_count = 2  # Reduced retries for speed
        for attempt in range(retry_count):
            try:
                url = f"https://www.linkedin.com/company/{company_name.lower().replace(' ', '-')}"
                
                # Resolve DNS first
                try:
                    result = await self.dns_resolver.query(url.split('/')[2], 'A')
                    ip = result[0].host
                except:
                    ip = None
                
                async with self.session.get(
                    url,
                    timeout=aiohttp.ClientTimeout(total=3),  # Aggressive timeout
                    ssl=False,  # Skip SSL verification
                    headers={'Host': url.split('/')[2]} if ip else {}
                ) as response:
                    if response.status == 200:
                        html = await response.text()
                        return await asyncio.get_event_loop().run_in_executor(
                            self.executor,
                            self._parse_linkedin_data,
                            BeautifulSoup(html, 'lxml')  # Using lxml for faster parsing
                        )
            except asyncio.TimeoutError:
                if attempt == retry_count - 1:
                    return {}
                await asyncio.sleep(0.1)
        return {}

    def _parse_linkedin_data(self, soup: BeautifulSoup) -> Dict:
        """Parse LinkedIn HTML with optimized selectors"""
        try:
            description = soup.find('section', {'class': 'description'})
            industry = soup.find('div', {'class': 'industry'})
            
            return {
                "description": description.get_text() if description else "",
                "industry": industry.get_text() if industry else ""
            }
        except:
            return {}

    async def _fetch_website_data(self, company_name: str) -> str:
        """Fetch website data with optimized timeout and error handling"""
        try:
            website = f"https://www.{company_name.lower().replace(' ', '')}.com"
            
            # Try to resolve DNS first
            try:
                result = await self.dns_resolver.query(website.split('/')[2], 'A')
                ip = result[0].host
            except:
                ip = None
            
            async with self.session.get(
                website,
                timeout=aiohttp.ClientTimeout(total=3),
                ssl=False,
                headers={'Host': website.split('/')[2]} if ip else {}
            ) as response:
                if response.status == 200:
                    html = await response.text()
                    return await asyncio.get_event_loop().run_in_executor(
                        self.executor,
                        self._parse_website_data,
                        html
                    )
        except:
            return ""
        return ""

    def _parse_website_data(self, html: str) -> str:
        """Parse website HTML efficiently"""
        try:
            soup = BeautifulSoup(html, 'lxml')
            paragraphs = soup.find_all('p')
            return "\n".join(p.get_text() for p in paragraphs)
        except:
            return ""
    
    async def _fetch_glassdoor_data(self, company_name: str) -> Dict:
        """Fetch company data from Glassdoor"""
        try:
            # Note: This is a mock implementation. In production, you'd use Glassdoor's API
            async with self.session.get(
                f"https://api.glassdoor.com/api/v1/companies",
                params={"name": company_name},
                headers={"Authorization": os.getenv("GLASSDOOR_API_KEY")},
                timeout=5
            ) as response:
                if response.status == 200:
                    data = await response.json()
                    return {
                        "culture": data.get("culture", ""),
                        "benefits": data.get("benefits", []),
                        "ratings": data.get("ratings", {})
                    }
        except:
            return {}
        return {}
        
    async def _fetch_news_data(self, company_name: str) -> Dict:
        """Fetch recent news about the company"""
        try:
            # Note: This is a mock implementation. In production, you'd use a news API
            async with self.session.get(
                "https://newsapi.org/v2/everything",
                params={
                    "q": company_name,
                    "sortBy": "publishedAt",
                    "pageSize": 5,
                    "apiKey": os.getenv("NEWS_API_KEY")
                },
                timeout=5
            ) as response:
                if response.status == 200:
                    data = await response.json()
                    return {
                        "articles": [
                            {
                                "title": article["title"],
                                "summary": article["description"],
                                "date": article["publishedAt"]
                            }
                            for article in data.get("articles", [])
                        ]
                    }
        except:
            return {"articles": []}
        return {"articles": []}
        
    async def _extract_keywords(self, text: str) -> List[str]:
        """Extract relevant keywords from text using AI"""
        try:
            response = await self.openai_client.chat.completions.create(
                messages=[{
                    "role": "user",
                    "content": f"""Extract 5-10 relevant keywords from this text that would be useful for a cover letter:
                    {text[:1000]}
                    
                    Return only the keywords, separated by commas."""
                }],
                model="deepseek-chat",
                temperature=0.3,
                max_tokens=100
            )
            
            keywords = response.choices[0].message.content.split(",")
            return [k.strip() for k in keywords if k.strip()]
        except:
            return []
            
    async def _analyze_company_culture(self, linkedin_desc: str, glassdoor_culture: str, website_content: str) -> Dict:
        """Analyze company culture from various sources"""
        try:
            combined_text = f"""
            LinkedIn: {linkedin_desc}
            Glassdoor: {glassdoor_culture}
            Website: {website_content[:500]}
            """
            
            response = await self.openai_client.chat.completions.create(
                messages=[{
                    "role": "user",
                    "content": f"""Analyze this company's culture and values based on the following sources:
                    {combined_text}
                    
                    Return a JSON object with:
                    1. A list of 3-5 core company values
                    2. A brief description of the work environment
                    Format: {{"values": [], "environment": ""}}"""
                }],
                model="deepseek-chat",
                temperature=0.3,
                max_tokens=200
            )
            
            return json.loads(response.choices[0].message.content)
        except:
            return {"values": [], "environment": ""}
            
    async def _extract_tech_stack(self, website_content: str, company_desc: str) -> List[str]:
        """Extract technology stack mentioned in company materials"""
        try:
            combined_text = f"""
            Description: {company_desc}
            Website: {website_content[:500]}
            """
            
            response = await self.openai_client.chat.completions.create(
                messages=[{
                    "role": "user",
                    "content": f"""Extract the technology stack and tools mentioned in this text:
                    {combined_text}
                    
                    Return only the technology names, separated by commas."""
                }],
                model="deepseek-chat",
                temperature=0.3,
                max_tokens=100
            )
            
            tech_stack = response.choices[0].message.content.split(",")
            return [tech.strip() for tech in tech_stack if tech.strip()]
        except:
            return []
    
    async def generate_single_letter(self, company_name: str, position: str, user_profile: Dict, 
                                   template_id: str, formats: List[str], notes: str = "", company_info: Dict = None) -> Dict:
        """Generate a single cover letter with specified formats"""
        try:
            # Generate cache key
            cache_key = f"letter:{company_name}:{position}:{template_id}:{hash(json.dumps(user_profile, default=self._serialize_date))}"
            
            # Check Redis cache
            if self.redis_client:
                cached_letter = await self.redis_client.get(cache_key)
                if cached_letter:
                    try:
                        cached_data = json.loads(cached_letter)
                        if all(os.path.exists(cached_data[k]) for k in cached_data.keys()):
                            return cached_data
                    except:
                        pass
            
            # Get template (should be cached from earlier)
            template = await self._get_template(template_id)
            
            # Generate content with retry
            content = await self._generate_content_with_retry(
                company_info=company_info or await self.gather_company_info(company_name),
                position=position,
                user_profile=user_profile,
                template=template,
                notes=notes
            )
            
            result = {}
            
            # Create documents based on selected formats
            if 'docx' in formats:
                docx_path = await self._create_docx(content, company_name, position)
                result['docx_path'] = docx_path
            
            if 'pdf' in formats:
                # If we have a DOCX, convert from it; otherwise create directly
                pdf_path = await self._create_pdf(
                    content, 
                    company_name, 
                    position,
                    docx_path=result.get('docx_path')
                )
                result['pdf_path'] = pdf_path
            
            # Cache the result
            await self._cache_result(cache_key, result)
            
            return result
            
        except Exception as e:
            raise Exception(f"Error generating letter for {company_name}: {str(e)}")

    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    async def _generate_content_with_retry(self, **kwargs):
        """Generate content with retry logic"""
        try:
            return await self._generate_content(**kwargs)
        except Exception as e:
            self.logger.error(f"Error in content generation attempt: {str(e)}")
            raise

    async def _get_template(self, template_id: str) -> Dict:
        """Get template with caching"""
        # First try memory cache
        if template_id in self.template_cache:
            return self.template_cache[template_id]
            
        # Then try Redis if available
        if self.redis_client:
            cached_template = await self._safe_redis_operation(
                lambda: self.redis_client.get(f"template:{template_id}")
            )
            if cached_template:
                template = json.loads(cached_template)
                self.template_cache[template_id] = template
                return template
        
        # Fall back to fetching fresh
        template = get_template_prompt(template_id)
        self.template_cache[template_id] = template
        return template

    async def _generate_content(self, company_info: Dict, position: str, user_profile: Dict, template: Dict, notes: str = "") -> str:
        """Generate cover letter content using AI"""
        try:
            # Prepare context for generation
            context = {
                "company": {
                    "name": company_info.get("name", ""),
                    "description": company_info.get("description", ""),
                    "culture": company_info.get("culture", {}),
                    "technologies": company_info.get("technologies", [])
                },
                "position": position,
                "candidate": {
                    "name": user_profile.get("name", ""),
                    "experience": user_profile.get("experience", [])[:3],  # Most recent 3 experiences
                    "skills": user_profile.get("skills", []),
                    "education": user_profile.get("education", []),
                    "summary": user_profile.get("summary", "")
                },
                "notes": notes
            }

            # Get template prompt
            prompt = template.get("prompt_template", "")
            if not prompt:
                raise ValueError("Template prompt not found")

            # Generate content with AI
            response = await self.openai_client.chat.completions.create(
                messages=[{
                    "role": "system",
                    "content": "You are an expert cover letter writer. Generate a professional, compelling cover letter."
                }, {
                    "role": "user",
                    "content": f"{prompt}\n\nContext: {json.dumps(context, default=self._serialize_date)}"
                }],
                model="deepseek-chat",
                temperature=0.7,
                max_tokens=1000
            )

            return response.choices[0].message.content

        except Exception as e:
            self.logger.error(f"Error generating content: {str(e)}")
            raise

    async def _create_docx(self, content: str, company_name: str, position: str) -> str:
        """Create DOCX file with improved formatting"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Set margins
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            def create_doc():
                try:
                    # Create company-specific directory
                    company_dir = self.output_dir / f"{company_name}"
                    company_dir.mkdir(parents=True, exist_ok=True)
                    
                    # Split content into paragraphs
                    paragraphs = content.split('\n\n')
                    
                    # Process each paragraph
                    for p in paragraphs:
                        if p.strip():
                            para = doc.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            
                            # Handle bold text (text between ** **)
                            parts = p.split('**')
                            for i, part in enumerate(parts):
                                if not part:  # Skip empty parts
                                    continue
                                    
                                run = para.add_run(part)
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(12)
                                
                                # Make every other part bold (parts between ** **)
                                if i % 2 == 1:  # Odd indices are bold
                                    run.bold = True
                    
                    # Save with a descriptive filename
                    safe_position = "".join(x for x in position if x.isalnum() or x in (' ', '-', '_')).strip()
                    filename = company_dir / f"{company_name}_{safe_position}.docx"
                    doc.save(str(filename))
                    return str(filename)
                except Exception as e:
                    self.logger.error(f"Error in DOCX creation: {str(e)}")
                    raise
            
            return await asyncio.get_event_loop().run_in_executor(
                self.executor,
                create_doc
            )
        except Exception as e:
            self.logger.error(f"Error creating DOCX: {str(e)}")
            raise

    async def _create_pdf(self, content: str, company_name: str, position: str, docx_path: str = None) -> str:
        """Create PDF file directly using reportlab"""
        try:
            # Create company-specific directory
            company_dir = self.output_dir / f"{company_name}"
            company_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate unique filename
            safe_position = "".join(x for x in position if x.isalnum() or x in (' ', '-', '_')).strip()
            filename = company_dir / f"{company_name}_{safe_position}.pdf"
            
            def create_pdf():
                try:
                    from reportlab.lib import colors
                    from reportlab.lib.pagesizes import letter
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.units import inch
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    
                    # Create document with proper margins
                    doc = SimpleDocTemplate(
                        str(filename),
                        pagesize=letter,
                        rightMargin=1*inch,
                        leftMargin=1*inch,
                        topMargin=1*inch,
                        bottomMargin=1*inch
                    )
                    
                    # Define styles
                    styles = getSampleStyleSheet()
                    normal_style = ParagraphStyle(
                        'CustomNormal',
                        parent=styles['Normal'],
                        fontSize=12,
                        leading=14,
                        fontName='Times-Roman'
                    )
                    
                    # Process content
                    story = []
                    paragraphs = content.split('\n\n')
                    
                    for p in paragraphs:
                        if p.strip():
                            # Handle bold text (between ** **)
                            parts = p.split('**')
                            formatted_text = ''
                            for i, part in enumerate(parts):
                                if i % 2 == 1:  # Odd indices are bold
                                    formatted_text += f'<b>{part}</b>'
                                else:
                                    formatted_text += part
                            
                            para = Paragraph(formatted_text, normal_style)
                            story.append(para)
                            story.append(Spacer(1, 12))
                    
                    # Build PDF
                    doc.build(story)
                    return str(filename)
                    
                except Exception as e:
                    self.logger.error(f"Error in PDF creation: {str(e)}")
                    raise
            
            # Create PDF in executor
            return await asyncio.get_event_loop().run_in_executor(
                self.executor,
                create_pdf
            )
            
        except Exception as e:
            self.logger.error(f"Error in PDF creation: {str(e)}")
            raise
    
    async def generate_preview(self, template_settings: dict, format_type: str = 'doc') -> str:
        """Generate a preview of the cover letter with the given template settings"""
        try:
            # Use a sample company and position for the preview
            sample_data = {
                "company_name": "Example Corp",
                "position": "Software Engineer",
                "notes": "Sample preview for customization"
            }
            
            # Create a mock company data instead of fetching
            company_data = {
                "name": sample_data["company_name"],
                "description": "A leading technology company focused on innovation",
                "industry": "Technology",
                "culture": {
                    "values": ["Innovation", "Excellence", "Teamwork"],
                    "work_environment": "Dynamic and collaborative workplace"
                }
            }
            
            # Generate content using the template settings
            content = await self._generate_content(
                company_info=company_data,
                position=sample_data["position"],
                user_profile={
                    "name": "John Doe",
                    "current_role": "Senior Developer",
                    "experience": [
                        {"title": "Lead Developer", "company": "Tech Co", "duration": "3 years"},
                        {"title": "Software Engineer", "company": "Startup Inc", "duration": "2 years"}
                    ],
                    "skills": ["Python", "JavaScript", "React", "Node.js", "AWS"]
                },
                template=template_settings,
                notes=""
            )
            
            # Format the content with proper HTML structure
            formatted_content = (
                '<div class="font-serif leading-relaxed">'
                f'{content.replace(chr(10), "<br>")}'
                '</div>'
            )
            
            if format_type == 'pdf':
                return await self.format_content_for_pdf(formatted_content)
            else:
                return formatted_content

        except Exception as e:
            self.logger.error(f"Error in generate_preview: {str(e)}", exc_info=True)
            raise
            
    def format_content_for_doc(self, content: str) -> str:
        """Format the content for DOC preview"""
        # Convert newlines to <br> tags and wrap paragraphs
        paragraphs = content.split('\n\n')
        formatted_paragraphs = []
        for p in paragraphs:
            if p.strip():
                # Replace newlines with <br> before using in f-string
                p_with_br = p.replace('\n', '<br>')
                formatted_paragraphs.append(f'<p class="mb-4">{p_with_br}</p>')
        return '\n'.join(formatted_paragraphs)
        
    async def format_content_for_pdf(self, content: str) -> str:
        """Format the content for PDF preview"""
        # Add any PDF-specific formatting
        formatted_content = content.replace('\n\n', '</p><p>').replace('\n', '<br>')
        
        # Build HTML content without f-strings
        html_parts = [
            '<html>',
            '<head>',
            '<style>',
            'body { font-family: "Times New Roman", serif; font-size: 12pt; line-height: 1.5; }',
            'p { margin-bottom: 1em; }',
            '</style>',
            '</head>',
            '<body>',
            formatted_content,
            '</body>',
            '</html>'
        ]
        
        return '\n'.join(html_parts)
        
    async def save_preview_pdf(self, content: str, output_path: str):
        """Save the preview content as a PDF file"""
        def create_pdf():
            import pypandoc
            pypandoc.convert_text(
                content,
                'pdf',
                format='html',
                outputfile=output_path,
                extra_args=['--pdf-engine=weasyprint']
            )
            
        await asyncio.get_event_loop().run_in_executor(self.executor, create_pdf)

    async def _package_results(self, results: List[Dict]) -> Dict:
        """Package the generated cover letter results into a single response"""
        try:
            # Create a zip file containing all generated documents
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            zip_filename = f"generated_letters_{timestamp}.zip"
            
            with zipfile.ZipFile(zip_filename, 'w') as zipf:
                for result in results:
                    if result and isinstance(result, dict):
                        # Add DOCX file if it exists
                        if 'docx_path' in result and os.path.exists(result['docx_path']):
                            zipf.write(result['docx_path'], os.path.basename(result['docx_path']))
                            await self._async_remove(result['docx_path'])
                            
                        # Add PDF file if it exists
                        if 'pdf_path' in result and os.path.exists(result['pdf_path']):
                            zipf.write(result['pdf_path'], os.path.basename(result['pdf_path']))
                            await self._async_remove(result['pdf_path'])
            
            return {
                'zip_file': zip_filename,
                'total_letters': len(results),
                'successful_generations': len([r for r in results if r and isinstance(r, dict)]),
                'timestamp': timestamp
            }
            
        except Exception as e:
            self.logger.error(f"Error packaging results: {str(e)}")
            raise

    def _serialize_date(self, obj):
        """Helper method to serialize dates"""
        if isinstance(obj, (datetime, date)):
            return obj.isoformat()
        return str(obj)

    def get_available_templates(self):
        """Get information about available templates."""
        from templates.cover_letter_templates import get_template_info
        return get_template_info() 

    async def _cache_result(self, cache_key: str, result: Dict):
        """Cache the result of a letter generation"""
        if self.redis_client:
            try:
                await self.redis_client.set(
                    cache_key,
                    json.dumps(result),
                    ex=3600  # Cache for 1 hour
                )
            except Exception as e:
                self.logger.warning(f"Failed to cache result: {str(e)}") 